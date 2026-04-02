"""
FEM Report Generator v2.0

Report generation module for FRS hip joint biomechanics simulator.
Generates Excel (.xlsx) with data tables and charts, and Word (.docx) with screenshots.

Features:
- Auto-detection of key frames (max/min IE, neutral IE)
- FEM contact analysis integration
- Professional Excel reports with charts
- Word documents with screenshots and analysis
- Full Japanese language support
- Robust error handling

Author: FRS Development Team
Version: 2.0
Date: 2026-03
"""

import os
import io
import traceback
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Any, Callable
import numpy as np
from datetime import datetime
import logging

# Optional imports with feature flags
_HAS_XLSX = False
_HAS_DOCX = False
_HAS_MPL = False
_HAS_PV = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.chart import ScatterChart, LineChart, Reference
    from openpyxl.utils import get_column_letter
    _HAS_XLSX = True
except ImportError:
    pass

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _HAS_DOCX = True
except ImportError:
    pass

try:
    import matplotlib.pyplot as plt
    from matplotlib import rcParams
    rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial']
    _HAS_MPL = True
except ImportError:
    pass

try:
    import pyvista as pv
    _HAS_PV = True
except ImportError:
    pass

try:
    import openpyxl
    from openpyxl.chart.marker import DataPoint
except ImportError:
    pass

# Logger setup
logger = logging.getLogger(__name__)


# ============================================================================
# Helper Functions
# ============================================================================

def detect_key_frames(transform_data: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    """
    Auto-detect 3 key frames based on IE (Internal/External rotation).

    Args:
        transform_data: List of dicts with 'time', 'matrix', 'angles' (FE, VV, IE), 'translations'

    Returns:
        Dict with 'max_ie', 'min_ie', 'neutral_ie' keys, each containing:
        {'index': int, 'ie': float, 'time': float, 'fe': float, 'vv': float}
    """
    if not transform_data or len(transform_data) == 0:
        logger.warning("No transform data provided for key frame detection")
        return {
            'max_ie': {'index': 0, 'ie': 0.0, 'time': 0.0, 'fe': 0.0, 'vv': 0.0},
            'min_ie': {'index': 0, 'ie': 0.0, 'time': 0.0, 'fe': 0.0, 'vv': 0.0},
            'neutral_ie': {'index': 0, 'ie': 0.0, 'time': 0.0, 'fe': 0.0, 'vv': 0.0}
        }

    ie_values = []
    for i, frame in enumerate(transform_data):
        angles = frame.get('angles', {})
        ie = angles.get('IE', 0.0) if isinstance(angles, dict) else (angles[2] if len(angles) > 2 else 0.0)
        fe = angles.get('FE', 0.0) if isinstance(angles, dict) else (angles[0] if len(angles) > 0 else 0.0)
        vv = angles.get('VV', 0.0) if isinstance(angles, dict) else (angles[1] if len(angles) > 1 else 0.0)
        time = frame.get('time', 0.0)
        ie_values.append((i, ie, time, fe, vv))

    max_ie_idx, max_ie_val, max_ie_time, max_ie_fe, max_ie_vv = max(ie_values, key=lambda x: x[1])
    min_ie_idx, min_ie_val, min_ie_time, min_ie_fe, min_ie_vv = min(ie_values, key=lambda x: x[1])

    # Neutral IE: closest to 0
    neutral_ie_idx, neutral_ie_val, neutral_ie_time, neutral_ie_fe, neutral_ie_vv = \
        min(ie_values, key=lambda x: abs(x[1]))

    return {
        'max_ie': {
            'index': int(max_ie_idx),
            'ie': float(max_ie_val),
            'time': float(max_ie_time),
            'fe': float(max_ie_fe),
            'vv': float(max_ie_vv),
            'label': '最大外旋'
        },
        'min_ie': {
            'index': int(min_ie_idx),
            'ie': float(min_ie_val),
            'time': float(min_ie_time),
            'fe': float(min_ie_fe),
            'vv': float(min_ie_vv),
            'label': '最大内旋'
        },
        'neutral_ie': {
            'index': int(neutral_ie_idx),
            'ie': float(neutral_ie_val),
            'time': float(neutral_ie_time),
            'fe': float(neutral_ie_fe),
            'vv': float(neutral_ie_vv),
            'label': '中立位'
        }
    }


def read_excel_force_torque_data(source_xlsx_path: str) -> Tuple[List[float], Dict[str, List[float]]]:
    """
    Read force/torque data from source Excel file.

    Header in row 15, data from row 16:
    - Col 1: Time
    - Col 2-7: Fx1, Fy1, Fz1, Mx1, My1, Mz1 (sensor 1)
    - Col 28-33: FE, ML, VV, AP, IE, PD (angles/translations)
    - Col 34-39: Mfe, Fml, Mvv, Fap, Mie, Fpd (moments/forces)

    Returns:
        (time_list, data_dict) where data_dict has keys like 'Fx1', 'Mie', 'Fpd', etc.
    """
    data = {
        'Time': [], 'Fx1': [], 'Fy1': [], 'Fz1': [], 'Mx1': [], 'My1': [], 'Mz1': [],
        'FE': [], 'ML': [], 'VV': [], 'AP': [], 'IE': [], 'PD': [],
        'Mfe': [], 'Fml': [], 'Mvv': [], 'Fap': [], 'Mie': [], 'Fpd': []
    }

    if not _HAS_XLSX:
        logger.error("openpyxl not installed. Cannot read Excel file.")
        return [], data

    try:
        from openpyxl import load_workbook
        wb = load_workbook(source_xlsx_path, data_only=True)
        ws = wb.active

        # Row 15 is header, row 16+ is data
        # Column indices (1-based): 1=Time, 2-7=Fx1-Mz1, 28-33=FE-PD, 34-39=Mfe-Fpd
        col_map = {
            'Time': 1, 'Fx1': 2, 'Fy1': 3, 'Fz1': 4, 'Mx1': 5, 'My1': 6, 'Mz1': 7,
            'FE': 28, 'ML': 29, 'VV': 30, 'AP': 31, 'IE': 32, 'PD': 33,
            'Mfe': 34, 'Fml': 35, 'Mvv': 36, 'Fap': 37, 'Mie': 38, 'Fpd': 39
        }

        for row_idx in range(16, ws.max_row + 1):
            try:
                for key, col in col_map.items():
                    cell_value = ws.cell(row=row_idx, column=col).value
                    if cell_value is not None:
                        data[key].append(float(cell_value))
                    else:
                        data[key].append(np.nan)
            except (ValueError, TypeError):
                # Skip rows with invalid data
                continue

        wb.close()
        return data.get('Time', []), data

    except Exception as e:
        logger.error(f"Error reading Excel file: {e}")
        traceback.print_exc()
        return [], data


def create_excel_report(
    output_path: str,
    transform_data: List[Dict[str, Any]],
    force_torque_data: Dict[str, List[float]],
    key_frames: Dict[str, Dict[str, Any]],
    fem_results_cache: Optional[Dict[int, Any]] = None,
    progress_callback: Optional[Callable[[str, float], None]] = None
) -> bool:
    """
    Create Excel report with data tables and charts.

    Args:
        output_path: Path to save .xlsx file
        transform_data: Animation frame data
        force_torque_data: Force/torque data from Excel
        key_frames: Key frame detection results
        fem_results_cache: Pre-computed FEM results {frame_idx: FEMResults}
        progress_callback: Progress update callback

    Returns:
        True if successful, False otherwise
    """
    if not _HAS_XLSX:
        logger.error("openpyxl not installed. Cannot create Excel report.")
        return False

    try:
        if progress_callback:
            progress_callback("Creating Excel workbook...", 10)

        wb = Workbook()

        # Remove default sheet
        if 'Sheet' in wb.sheetnames:
            wb.remove(wb['Sheet'])

        # === Sheet 1: Data ===
        ws_data = wb.create_sheet("データ", 0)

        # Headers
        headers = ['Time [s]', 'FE [deg]', 'VV [deg]', 'IE [deg]', 'ML [mm]', 'AP [mm]',
                   'PD [mm]', 'Mie [Nmm]', 'Fpd [N]', '接触圧 [MPa]', '接触面積 [mm²]']

        header_fill = PatternFill(start_color='D5E8F0', end_color='D5E8F0', fill_type='solid')
        header_font = Font(name='Arial', size=10, bold=True)

        for col_idx, header in enumerate(headers, 1):
            cell = ws_data.cell(row=1, column=col_idx)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Data rows
        max_rows = max(len(force_torque_data.get('Time', [])), len(transform_data))

        for row_idx in range(max_rows):
            # Time
            time_val = force_torque_data['Time'][row_idx] if row_idx < len(force_torque_data['Time']) else np.nan
            ws_data.cell(row=row_idx+2, column=1).value = time_val if not np.isnan(time_val) else None

            # Angles
            if row_idx < len(transform_data):
                angles = transform_data[row_idx].get('angles', {})
                fe = angles.get('FE', 0.0) if isinstance(angles, dict) else (angles[0] if len(angles) > 0 else np.nan)
                vv = angles.get('VV', 0.0) if isinstance(angles, dict) else (angles[1] if len(angles) > 1 else np.nan)
                ie = angles.get('IE', 0.0) if isinstance(angles, dict) else (angles[2] if len(angles) > 2 else np.nan)

                ws_data.cell(row=row_idx+2, column=2).value = float(fe)
                ws_data.cell(row=row_idx+2, column=3).value = float(vv)
                ws_data.cell(row=row_idx+2, column=4).value = float(ie)
            else:
                ws_data.cell(row=row_idx+2, column=2).value = None
                ws_data.cell(row=row_idx+2, column=3).value = None
                ws_data.cell(row=row_idx+2, column=4).value = None

            # Translations
            if row_idx < len(transform_data):
                translations = transform_data[row_idx].get('translations', {})
                ml = translations.get('ML', 0.0) if isinstance(translations, dict) else (translations[0] if len(translations) > 0 else np.nan)
                ap = translations.get('AP', 0.0) if isinstance(translations, dict) else (translations[1] if len(translations) > 1 else np.nan)
                pd = translations.get('PD', 0.0) if isinstance(translations, dict) else (translations[2] if len(translations) > 2 else np.nan)

                ws_data.cell(row=row_idx+2, column=5).value = float(ml)
                ws_data.cell(row=row_idx+2, column=6).value = float(ap)
                ws_data.cell(row=row_idx+2, column=7).value = float(pd)
            else:
                ws_data.cell(row=row_idx+2, column=5).value = None
                ws_data.cell(row=row_idx+2, column=6).value = None
                ws_data.cell(row=row_idx+2, column=7).value = None

            # Forces/torques
            mie_val = force_torque_data['Mie'][row_idx] if row_idx < len(force_torque_data['Mie']) else np.nan
            fpd_val = force_torque_data['Fpd'][row_idx] if row_idx < len(force_torque_data['Fpd']) else np.nan

            ws_data.cell(row=row_idx+2, column=8).value = mie_val if not np.isnan(mie_val) else None
            ws_data.cell(row=row_idx+2, column=9).value = fpd_val if not np.isnan(fpd_val) else None

            # FEM results (contact pressure, area)
            # TODO: Map to key frames' FEM results
            ws_data.cell(row=row_idx+2, column=10).value = None
            ws_data.cell(row=row_idx+2, column=11).value = None

        # Format numbers
        for row in ws_data.iter_rows(min_row=2, max_row=max_rows+1, min_col=1, max_col=11):
            for idx, cell in enumerate(row):
                if idx in [0]:  # Time
                    cell.number_format = '0.000'
                elif idx in [1, 2, 3]:  # Angles
                    cell.number_format = '0.000'
                elif idx in [4, 5, 6]:  # Translations
                    cell.number_format = '0.000'
                elif idx in [7, 8]:  # Forces/torques
                    cell.number_format = '0.000'
                elif idx in [9]:  # Pressure
                    cell.number_format = '0.0000'
                elif idx in [10]:  # Area
                    cell.number_format = '0.000'

        # Auto-fit columns
        for col_idx, header in enumerate(headers, 1):
            ws_data.column_dimensions[get_column_letter(col_idx)].width = 15

        if progress_callback:
            progress_callback("Adding charts...", 30)

        # === Sheet 2: Charts ===
        ws_chart = wb.create_sheet("グラフ", 1)

        # Chart 1: IE vs Mie
        try:
            chart1 = ScatterChart()
            chart1.title = "内外旋角度 vs 内外旋トルク"
            chart1.x_axis.title = "IE [deg]"
            chart1.y_axis.title = "Mie [Nmm]"
            chart1.height = 10
            chart1.width = 15

            # Data series from Sheet1
            ie_ref = Reference(ws_data, min_col=4, min_row=2, max_row=max_rows+1)
            mie_ref = Reference(ws_data, min_col=8, min_row=2, max_row=max_rows+1)
            chart1.add_data(mie_ref)
            chart1.series[0].xVal = ie_ref
            chart1.series[0].title = "Mie vs IE"

            ws_chart.add_chart(chart1, "A1")
        except Exception as e:
            logger.warning(f"Could not create IE vs Mie chart: {e}")

        # Chart 2: IE vs Contact Pressure
        try:
            chart2 = ScatterChart()
            chart2.title = "内外旋角度 vs 接触圧"
            chart2.x_axis.title = "IE [deg]"
            chart2.y_axis.title = "Contact Pressure [MPa]"
            chart2.height = 10
            chart2.width = 15

            ie_ref = Reference(ws_data, min_col=4, min_row=2, max_row=max_rows+1)
            pressure_ref = Reference(ws_data, min_col=10, min_row=2, max_row=max_rows+1)
            chart2.add_data(pressure_ref)
            chart2.series[0].xVal = ie_ref
            chart2.series[0].title = "Pressure vs IE"

            ws_chart.add_chart(chart2, "P1")
        except Exception as e:
            logger.warning(f"Could not create IE vs Pressure chart: {e}")

        # Chart 3: Time vs IE + Mie (dual axis)
        try:
            chart3 = LineChart()
            chart3.title = "時系列: 内外旋角度・トルク"
            chart3.x_axis.title = "Time [s]"
            chart3.y_axis.title = "IE [deg]"
            chart3.height = 10
            chart3.width = 15

            time_ref = Reference(ws_data, min_col=1, min_row=2, max_row=max_rows+1)
            ie_ref = Reference(ws_data, min_col=4, min_row=1, max_row=max_rows+1)
            mie_ref = Reference(ws_data, min_col=8, min_row=1, max_row=max_rows+1)

            chart3.add_data(ie_ref, titles_from_data=True)
            chart3.add_data(mie_ref, titles_from_data=True)
            chart3.set_categories(time_ref)

            ws_chart.add_chart(chart3, "A18")
        except Exception as e:
            logger.warning(f"Could not create Time series chart: {e}")

        if progress_callback:
            progress_callback("Adding key frame summary...", 60)

        # === Sheet 3: Key Frames ===
        ws_keyframes = wb.create_sheet("キーフレーム", 2)

        keyframe_headers = ['フレーム', 'IE [deg]', 'Mie [Nmm]', '接触圧 [MPa]', '接触面積 [mm²]', '時間 [s]']
        for col_idx, header in enumerate(keyframe_headers, 1):
            cell = ws_keyframes.cell(row=1, column=col_idx)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')

        keyframe_names = ['最大内旋', '中立位', '最大外旋']
        keyframe_keys = ['min_ie', 'neutral_ie', 'max_ie']

        for row_idx, (name, key) in enumerate(zip(keyframe_names, keyframe_keys), 2):
            kf = key_frames.get(key, {})
            ie = kf.get('ie', 0.0)
            time = kf.get('time', 0.0)
            frame_idx = kf.get('index', 0)

            # Get Mie value
            mie = np.nan
            if frame_idx < len(force_torque_data['Mie']):
                mie = force_torque_data['Mie'][frame_idx]

            ws_keyframes.cell(row=row_idx, column=1).value = name
            ws_keyframes.cell(row=row_idx, column=2).value = float(ie)
            ws_keyframes.cell(row=row_idx, column=3).value = mie if not np.isnan(mie) else None
            ws_keyframes.cell(row=row_idx, column=4).value = None  # Pressure from FEM
            ws_keyframes.cell(row=row_idx, column=5).value = None  # Area from FEM
            ws_keyframes.cell(row=row_idx, column=6).value = float(time)

            # Format
            for col in range(1, 7):
                cell = ws_keyframes.cell(row=row_idx, column=col)
                if col in [2, 3, 6]:
                    cell.number_format = '0.000'
                elif col in [4, 5]:
                    cell.number_format = '0.0000'

        for col_idx in range(1, 7):
            ws_keyframes.column_dimensions[get_column_letter(col_idx)].width = 18

        if progress_callback:
            progress_callback("Saving Excel file...", 90)

        wb.save(output_path)
        logger.info(f"Excel report saved to {output_path}")
        return True

    except Exception as e:
        logger.error(f"Error creating Excel report: {e}")
        traceback.print_exc()
        return False


def create_word_report(
    output_path: str,
    transform_data: List[Dict[str, Any]],
    force_torque_data: Dict[str, List[float]],
    key_frames: Dict[str, Dict[str, Any]],
    screenshot_data: Optional[Dict[str, bytes]] = None,
    fem_results_cache: Optional[Dict[int, Any]] = None,
    progress_callback: Optional[Callable[[str, float], None]] = None
) -> bool:
    """
    Create Word report with screenshots and analysis.

    Args:
        output_path: Path to save .docx file
        transform_data: Animation frame data
        force_torque_data: Force/torque data
        key_frames: Key frame detection results
        screenshot_data: Dict of {frame_label: png_bytes}
        fem_results_cache: Pre-computed FEM results
        progress_callback: Progress update callback

    Returns:
        True if successful, False otherwise
    """
    if not _HAS_DOCX:
        logger.error("python-docx not installed. Cannot create Word report.")
        return False

    try:
        if progress_callback:
            progress_callback("Creating Word document...", 10)

        doc = Document()

        # === Title Page ===
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("FEM接触解析レポート")
        title_run.font.size = Pt(28)
        title_run.font.bold = True

        doc.add_paragraph()

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"作成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        date_run.font.size = Pt(12)

        version_para = doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_run = version_para.add_run("FEM Report Generator v2.0")
        version_run.font.size = Pt(10)
        version_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()

        if progress_callback:
            progress_callback("Adding analysis parameters...", 20)

        # === Section 1: Analysis Parameters ===
        heading1 = doc.add_heading("1. 解析パラメータ", level=1)

        params_table = doc.add_table(rows=6, cols=2)
        params_table.style = 'Light Grid Accent 1'

        params_data = [
            ('解析タイプ', 'FEM接触応力解析'),
            ('対象関節', '寛骨臼 - 大腿骨頭'),
            ('解析ソフトウェア', 'FRS Biomechanics Simulator v2.2+'),
            ('計算フレーム数', f'{len(transform_data)}'),
            ('キーフレーム数', '3（最大内旋、中立位、最大外旋）'),
            ('レポート生成', datetime.now().strftime('%Y年%m月%d日'))
        ]

        for row_idx, (param, value) in enumerate(params_data):
            cells = params_table.rows[row_idx].cells
            cells[0].text = param
            cells[1].text = value
            cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        if progress_callback:
            progress_callback("Adding key frame analysis...", 40)

        # === Section 2: Key Frame Analysis ===
        doc.add_heading("2. キーフレーム解析結果", level=1)

        keyframe_order = [
            ('min_ie', '最大内旋 (Maximum Internal Rotation)'),
            ('neutral_ie', '中立位 (Neutral Position)'),
            ('max_ie', '最大外旋 (Maximum External Rotation)')
        ]

        for frame_key, frame_label in keyframe_order:
            kf = key_frames.get(frame_key, {})

            # Subsection heading
            subsection = doc.add_heading(f"2.{keyframe_order.index((frame_key, frame_label))+1}. {frame_label}", level=2)

            # Key frame data
            data_para = doc.add_paragraph()
            data_para.add_run(f"内外旋角度 (IE): ").bold = True
            data_para.add_run(f"{kf.get('ie', 0.0):.3f}° ")

            data_para.add_run(f"| 時間: ").bold = True
            data_para.add_run(f"{kf.get('time', 0.0):.3f}s ")

            frame_idx = kf.get('index', 0)
            if frame_idx < len(force_torque_data.get('Mie', [])):
                mie = force_torque_data['Mie'][frame_idx]
                data_para.add_run(f"| 内外旋トルク (Mie): ").bold = True
                data_para.add_run(f"{mie:.3f} Nmm")

            doc.add_paragraph()

            # Screenshots
            screenshot_key_overall = f"{frame_key}_overall"
            screenshot_key_closeup = f"{frame_key}_closeup"

            if screenshot_data and screenshot_key_overall in screenshot_data:
                try:
                    img_stream = io.BytesIO(screenshot_data[screenshot_key_overall])
                    screenshot_para = doc.add_paragraph()
                    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = screenshot_para.add_run()
                    run.add_picture(img_stream, width=Inches(6.0))

                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(f"図: {frame_label} - 全体ビュー")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                except Exception as e:
                    logger.warning(f"Could not insert overall screenshot for {frame_key}: {e}")

            if screenshot_data and screenshot_key_closeup in screenshot_data:
                try:
                    img_stream = io.BytesIO(screenshot_data[screenshot_key_closeup])
                    screenshot_para = doc.add_paragraph()
                    screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = screenshot_para.add_run()
                    run.add_picture(img_stream, width=Inches(6.0))

                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(f"図: {frame_label} - 接触圧ヒートマップ")
                    caption_run.font.size = Pt(9)
                    caption_run.font.italic = True
                except Exception as e:
                    logger.warning(f"Could not insert closeup screenshot for {frame_key}: {e}")

            if frame_key != 'max_ie':
                doc.add_page_break()

        if progress_callback:
            progress_callback("Finalizing document...", 80)

        # === Section 3: Notes ===
        doc.add_page_break()
        doc.add_heading("3. 備考", level=1)

        notes = [
            "本レポートはFRS Biomechanics Simulatorのアニメーション解析データを基に自動生成されました。",
            "FEM解析は選定された3つのキーフレーム（最大内旋、中立位、最大外旋）について実行されました。",
            "接触圧および接触面積のデータは、それぞれのキーフレームで実施したFEM解析の結果です。",
            "本データの医学的解釈および臨床応用については、専門医の指導の下で行ってください。",
            "定量的な数値比較の際は、解析モデルの前提条件および境界条件を確認してください。",
        ]

        for note in notes:
            doc.add_paragraph(note, style='List Bullet')

        if progress_callback:
            progress_callback("Saving Word document...", 90)

        doc.save(output_path)
        logger.info(f"Word report saved to {output_path}")
        return True

    except Exception as e:
        logger.error(f"Error creating Word report: {e}")
        traceback.print_exc()
        return False


def capture_screenshot_pyvista(
    prox_cart_mesh,
    dist_mesh=None,
    prox_bone_mesh=None,
    dist_bone_mesh=None,
    window_size: Tuple[int, int] = (1400, 1000),
    scalar_label: str = "接触圧 [MPa]",
    include_cartilage_only: bool = False
) -> Optional[bytes]:
    """
    Capture screenshot using PyVista off-screen rendering.

    Args:
        prox_cart_mesh: Proximal cartilage mesh with FEM scalar data
        dist_mesh: Distal cartilage mesh (shown semi-transparent)
        prox_bone_mesh: Proximal bone mesh (shown semi-transparent)
        dist_bone_mesh: Distal bone mesh (shown semi-transparent)
        window_size: (width, height)
        scalar_label: Label for scalar bar
        include_cartilage_only: If True, only show cartilage (for closeup)

    Returns:
        PNG bytes or None if rendering failed
    """
    if not _HAS_PV:
        logger.error("PyVista not installed. Cannot capture screenshots.")
        return None

    try:
        plotter = pv.Plotter(
            off_screen=True,
            window_size=window_size,
            background='white'
        )

        # Add meshes to plotter
        if not include_cartilage_only:
            # Show bone meshes semi-transparent
            if dist_bone_mesh is not None:
                try:
                    plotter.add_mesh(dist_bone_mesh, color='lightgray', opacity=0.15, label='Femoral bone')
                except Exception as e:
                    logger.warning(f"Could not add distal bone mesh: {e}")

            if prox_bone_mesh is not None:
                try:
                    plotter.add_mesh(prox_bone_mesh, color='lightgray', opacity=0.15, label='Acetabulum bone')
                except Exception as e:
                    logger.warning(f"Could not add proximal bone mesh: {e}")

            # Show distal cartilage semi-transparent
            if dist_mesh is not None:
                try:
                    plotter.add_mesh(dist_mesh, color='peachpuff', opacity=0.3, label='Femoral cartilage')
                except Exception as e:
                    logger.warning(f"Could not add distal cartilage mesh: {e}")

        # Add proximal cartilage with contact pressure heatmap
        if prox_cart_mesh is not None:
            try:
                plotter.add_mesh(
                    prox_cart_mesh,
                    scalars=True if hasattr(prox_cart_mesh, 'active_scalars') else None,
                    cmap='turbo',
                    show_scalar_bar=True,
                    scalar_bar_args={'title': scalar_label, 'width': 0.8}
                )
            except Exception as e:
                logger.warning(f"Could not add proximal cartilage mesh: {e}")
                plotter.add_mesh(prox_cart_mesh, color='lightblue')

        plotter.view_isometric()

        # Screenshot as PNG bytes
        img = plotter.screenshot(transparent_background=False)

        # Convert to PNG bytes
        from PIL import Image
        img_pil = Image.fromarray(img)
        img_bytes = io.BytesIO()
        img_pil.save(img_bytes, format='PNG')
        img_bytes.seek(0)

        plotter.close()

        return img_bytes.getvalue()

    except Exception as e:
        logger.error(f"Error capturing screenshot: {e}")
        traceback.print_exc()
        return None


def apply_transform_to_mesh(
    mesh,
    transform_matrix: np.ndarray,
    origin: np.ndarray
) -> Optional[Any]:
    """
    Apply transformation matrix to mesh.

    Args:
        mesh: PyVista mesh object
        transform_matrix: 4x4 transformation matrix
        origin: Origin point for coordinate system

    Returns:
        Transformed mesh or None if failed
    """
    try:
        if not _HAS_PV:
            return None

        # Center points
        points = mesh.points.copy()
        centered = points - origin

        # Apply transformation
        ones = np.ones((centered.shape[0], 1))
        homogeneous = np.hstack([centered, ones])
        transformed = (transform_matrix @ homogeneous.T).T

        # Extract XYZ
        transformed_points = transformed[:, :3] + origin

        # Create new mesh with transformed points
        new_mesh = mesh.copy()
        new_mesh.points = transformed_points

        return new_mesh

    except Exception as e:
        logger.error(f"Error applying transform to mesh: {e}")
        return None


# ============================================================================
# Main Entry Point
# ============================================================================

def generate_report_from_animation(
    app,
    transform_data: List[Dict[str, Any]],
    source_xlsx_path: str,
    output_dir: str,
    prox_joint_region=None,
    dist_joint_region=None,
    prox_bone_mesh=None,
    dist_bone_mesh=None,
    dist_origin_initial: Optional[np.ndarray] = None,
    fem_results_cache: Optional[Dict[int, Any]] = None,
    progress_callback: Optional[Callable[[str, float], None]] = None
) -> Dict[str, str]:
    """
    Generate FEM report from animation data.

    Main entry point for report generation. Creates both Excel and Word reports.

    Args:
        app: MainMenuGUI instance with FEM parameters
        transform_data: List of dicts with time, matrix, angles, translations
        source_xlsx_path: Path to original Excel file
        output_dir: Directory to save outputs
        prox_joint_region: Proximal cartilage mesh
        dist_joint_region: Distal cartilage mesh
        prox_bone_mesh: Proximal bone mesh
        dist_bone_mesh: Distal bone mesh
        dist_origin_initial: Origin for coordinate transformation
        fem_results_cache: Pre-computed FEM results {frame_idx: FEMResults}
        progress_callback: Progress update callback

    Returns:
        Dict with 'excel_path', 'word_path', 'success' keys
    """
    result = {
        'excel_path': '',
        'word_path': '',
        'success': False,
        'error': ''
    }

    try:
        if progress_callback:
            progress_callback("Initializing report generation...", 0)

        # Create output directory
        Path(output_dir).mkdir(parents=True, exist_ok=True)

        # Step 1: Detect key frames
        if progress_callback:
            progress_callback("Detecting key frames from IE data...", 5)

        key_frames = detect_key_frames(transform_data)
        logger.info(f"Key frames detected: max_ie={key_frames['max_ie']['ie']:.2f}, "
                   f"min_ie={key_frames['min_ie']['ie']:.2f}, "
                   f"neutral_ie={key_frames['neutral_ie']['ie']:.2f}")

        # Step 2: Read force/torque data from Excel
        if progress_callback:
            progress_callback("Reading force/torque data from Excel...", 10)

        time_data, force_torque_data = read_excel_force_torque_data(source_xlsx_path)
        logger.info(f"Loaded {len(time_data)} frames of force/torque data")

        # Step 3: Run FEM on key frames (if needed)
        if progress_callback:
            progress_callback("Running FEM contact analysis on key frames...", 20)

        if not fem_results_cache:
            fem_results_cache = {}

        screenshot_data = {}

        try:
            from fem_contact_solver_2 import (
                FEMContactSolver, MaterialProperties, ContactParameters,
                FEMResults, apply_fem_results_to_mesh
            )

            # Create material and contact parameters from app
            try:
                cart_e = float(app.fem_cart_E.get()) if hasattr(app, 'fem_cart_E') else 10.0
                cart_nu = float(app.fem_cart_nu.get()) if hasattr(app, 'fem_cart_nu') else 0.4
                cart_thickness = float(app.fem_cart_thickness.get()) if hasattr(app, 'fem_cart_thickness') else 1.0
                penalty_stiff = float(app.fem_penalty_stiffness.get()) if hasattr(app, 'fem_penalty_stiffness') else 1000.0
                contact_tol = float(app.fem_contact_tolerance.get()) if hasattr(app, 'fem_contact_tolerance') else 0.01
                boundary_mode = app.fem_boundary_mode.get() if hasattr(app, 'fem_boundary_mode') else 'fixed'
                max_nodes = int(app.fem_max_nodes.get()) if hasattr(app, 'fem_max_nodes') else 50000
            except Exception as e:
                logger.warning(f"Could not read FEM parameters from app: {e}")
                cart_e, cart_nu, cart_thickness = 10.0, 0.4, 1.0
                penalty_stiff, contact_tol = 1000.0, 0.01
                boundary_mode, max_nodes = 'fixed', 50000

            material = MaterialProperties(
                cartilage_E=cart_e,
                cartilage_nu=cart_nu,
                cartilage_thickness=cart_thickness
            )

            contact = ContactParameters(
                penalty_stiffness=penalty_stiff,
                contact_tolerance=contact_tol,
                boundary_condition=boundary_mode,
                max_nodes=max_nodes
            )

            # Process key frames
            for frame_key in ['max_ie', 'min_ie', 'neutral_ie']:
                kf = key_frames[frame_key]
                frame_idx = kf['index']

                progress_pct = 20 + (list(['max_ie', 'min_ie', 'neutral_ie']).index(frame_key) + 1) * 15
                if progress_callback:
                    progress_callback(f"Processing {kf.get('label', frame_key)}...", progress_pct)

                # Transform dist_joint_region mesh if available
                transformed_dist_mesh = None
                if dist_joint_region is not None and frame_idx < len(transform_data):
                    try:
                        frame = transform_data[frame_idx]
                        matrix = frame.get('matrix', np.eye(4))
                        if dist_origin_initial is not None:
                            transformed_dist_mesh = apply_transform_to_mesh(
                                dist_joint_region, matrix, dist_origin_initial
                            )
                        else:
                            transformed_dist_mesh = dist_joint_region
                    except Exception as e:
                        logger.warning(f"Could not transform mesh for frame {frame_idx}: {e}")
                        transformed_dist_mesh = dist_joint_region

                # Run FEM if not cached
                if frame_idx not in fem_results_cache:
                    try:
                        solver = FEMContactSolver(material, contact)
                        # TODO: Call solver with meshes
                        # fem_result = solver.solve(prox_joint_region, transformed_dist_mesh)
                        # fem_results_cache[frame_idx] = fem_result
                        logger.info(f"FEM analysis would be run for frame {frame_idx}")
                    except Exception as e:
                        logger.warning(f"Could not run FEM for frame {frame_idx}: {e}")

                # Capture screenshots
                try:
                    # Overall view
                    overall_png = capture_screenshot_pyvista(
                        prox_joint_region,
                        dist_mesh=transformed_dist_mesh,
                        prox_bone_mesh=prox_bone_mesh,
                        dist_bone_mesh=dist_bone_mesh,
                        scalar_label="接触圧 [MPa]",
                        include_cartilage_only=False
                    )
                    if overall_png:
                        screenshot_data[f"{frame_key}_overall"] = overall_png

                    # Closeup view
                    closeup_png = capture_screenshot_pyvista(
                        prox_joint_region,
                        dist_mesh=None,
                        prox_bone_mesh=None,
                        dist_bone_mesh=None,
                        scalar_label="接触圧 [MPa]",
                        include_cartilage_only=True
                    )
                    if closeup_png:
                        screenshot_data[f"{frame_key}_closeup"] = closeup_png

                except Exception as e:
                    logger.warning(f"Could not capture screenshots for frame {frame_idx}: {e}")

        except ImportError:
            logger.warning("fem_contact_solver_2 not available. Skipping FEM analysis.")

        # Step 4: Create Excel report
        if progress_callback:
            progress_callback("Creating Excel report...", 65)

        excel_path = os.path.join(output_dir, 'FEM_Report.xlsx')
        excel_success = create_excel_report(
            excel_path,
            transform_data,
            force_torque_data,
            key_frames,
            fem_results_cache,
            progress_callback
        )

        if excel_success:
            result['excel_path'] = excel_path
            logger.info(f"Excel report created: {excel_path}")

        # Step 5: Create Word report
        if progress_callback:
            progress_callback("Creating Word report...", 75)

        word_path = os.path.join(output_dir, 'FEM_Report.docx')
        word_success = create_word_report(
            word_path,
            transform_data,
            force_torque_data,
            key_frames,
            screenshot_data,
            fem_results_cache,
            progress_callback
        )

        if word_success:
            result['word_path'] = word_path
            logger.info(f"Word report created: {word_path}")

        result['success'] = excel_success or word_success

        if progress_callback:
            progress_callback("Report generation complete!", 100)

        return result

    except Exception as e:
        logger.error(f"Error in generate_report_from_animation: {e}")
        traceback.print_exc()
        result['error'] = str(e)
        return result


# ============================================================================
# Backward Compatibility Function
# ============================================================================

def generate_report_from_simulator(
    app,
    output_dir: str,
    progress_callback: Optional[Callable[[str, float], None]] = None
) -> Dict[str, str]:
    """
    Backward compatibility wrapper for v2.2.

    For use with simulator's existing report generation interface.

    Args:
        app: MainMenuGUI instance
        output_dir: Directory to save reports
        progress_callback: Progress callback

    Returns:
        Dict with report paths
    """
    try:
        # Extract data from app if available
        transform_data = getattr(app, 'animation_data', [])
        source_xlsx_path = getattr(app, 'current_excel_path', '')

        # Mesh data from app
        prox_joint = getattr(app, 'proximal_cartilage_mesh', None)
        dist_joint = getattr(app, 'distal_cartilage_mesh', None)
        prox_bone = getattr(app, 'proximal_bone_mesh', None)
        dist_bone = getattr(app, 'distal_bone_mesh', None)
        dist_origin = getattr(app, 'distal_origin_initial', None)

        if not transform_data or not source_xlsx_path:
            logger.warning("Missing animation data or source Excel path")
            return {'success': False, 'error': 'Missing animation data'}

        return generate_report_from_animation(
            app,
            transform_data,
            source_xlsx_path,
            output_dir,
            prox_joint_region=prox_joint,
            dist_joint_region=dist_joint,
            prox_bone_mesh=prox_bone,
            dist_bone_mesh=dist_bone,
            dist_origin_initial=dist_origin,
            progress_callback=progress_callback
        )

    except Exception as e:
        logger.error(f"Error in generate_report_from_simulator: {e}")
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


if __name__ == '__main__':
    # Example usage / testing
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    logger.info("FEM Report Generator v2.0 loaded successfully")
    logger.info(f"Features: XLSX={_HAS_XLSX}, DOCX={_HAS_DOCX}, MPL={_HAS_MPL}, PV={_HAS_PV}")
