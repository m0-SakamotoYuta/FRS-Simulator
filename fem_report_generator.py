# ============================================================================
# FEM Report Generator — FEM解析レポート自動生成モジュール
# ============================================================================
# FRS_Simulator 連携モジュール
#
# ■ 機能:
#   1. 解析結果サマリー（数値表）
#   2. ヒートマップのスクリーンショット（PyVista off-screen rendering）
#   3. 時系列グラフ（matplotlib）
#   4. パラメータ一覧
#
# ■ 出力: Word (.docx)
# ■ 言語: 日本語
#
# ■ 依存: python-docx, matplotlib, numpy, pyvista（オプション）
# ============================================================================

import os
import io
import tempfile
import datetime
from typing import Optional, List, Dict, Any, Tuple

import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    print("[警告] python-docxがインストールされていません。pip install python-docx")

try:
    import matplotlib
    matplotlib.use('Agg')  # non-interactive backend
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    _HAS_MPL = True
except ImportError:
    _HAS_MPL = False

try:
    import pyvista as pv
    _HAS_PV = True
except ImportError:
    _HAS_PV = False


# ============================================================================
# ヘルパー関数
# ============================================================================

def _find_japanese_font():
    """利用可能な日本語フォントを探す"""
    preferred = [
        'IPAexGothic', 'IPAGothic', 'Noto Sans CJK JP', 'Noto Sans JP',
        'Yu Gothic', 'MS Gothic', 'Meiryo', 'Hiragino Sans',
        'TakaoPGothic', 'VL PGothic',
    ]
    available = {f.name for f in fm.fontManager.ttflist}
    for font_name in preferred:
        if font_name in available:
            return font_name
    # フォールバック: CJKフォントを探す
    for f in fm.fontManager.ttflist:
        if any(kw in f.name.lower() for kw in ['gothic', 'cjk', 'japanese', 'jp']):
            return f.name
    return None


def _setup_matplotlib_japanese():
    """matplotlibの日本語表示設定"""
    if not _HAS_MPL:
        return
    font = _find_japanese_font()
    if font:
        plt.rcParams['font.family'] = font
    plt.rcParams['axes.unicode_minus'] = False


def _set_cell_shading(cell, color_hex: str):
    """セルの背景色を設定"""
    shading = cell._tc.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def _set_table_style(table):
    """テーブルの枠線スタイルを設定"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '999999',
        })
        borders.append(elem)
    tblPr.append(borders)


# ============================================================================
# スクリーンショット生成
# ============================================================================

def capture_heatmap_screenshot(
    prox_mesh,
    scalar_name: str = 'contact_pressure',
    cmap: str = 'jet',
    dist_mesh=None,
    camera_position: str = 'xy',
    window_size: Tuple[int, int] = (1200, 900),
    title: str = '',
) -> Optional[bytes]:
    """PyVistaでヒートマップのスクリーンショットをキャプチャ

    Args:
        prox_mesh: スカラーデータ付きPyVistaメッシュ
        scalar_name: 表示するスカラー名
        cmap: カラーマップ
        dist_mesh: 遠位メッシュ（半透明表示、任意）
        camera_position: カメラ位置
        window_size: 画像サイズ
        title: 画像タイトル

    Returns:
        PNG画像のバイト列、失敗時はNone
    """
    if not _HAS_PV:
        return None

    try:
        plotter = pv.Plotter(off_screen=True, window_size=window_size)
        plotter.set_background('white')

        scalar_labels = {
            'contact_pressure': '接触圧 [MPa]',
            'von_mises_stress': 'von Mises応力 [MPa]',
            'max_principal_stress': '最大主応力 [MPa]',
            'displacement_magnitude': '変位量 [mm]',
            'penetration_depth': '侵入量 [mm]',
        }
        label = scalar_labels.get(scalar_name, scalar_name)

        if scalar_name in prox_mesh.array_names:
            plotter.add_mesh(
                prox_mesh, scalars=scalar_name, cmap=cmap,
                show_edges=False, scalar_bar_args={'title': label, 'color': 'black'},
                opacity=1.0,
            )
        else:
            plotter.add_mesh(prox_mesh, color='lightblue', opacity=1.0)

        if dist_mesh is not None:
            plotter.add_mesh(dist_mesh, color='#FFB6C1', opacity=0.3, show_edges=False)

        plotter.add_axes()
        if title:
            plotter.add_text(title, position='upper_left', font_size=14, color='black')

        plotter.camera_position = camera_position

        # スクリーンショットをバイト列として取得
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            tmp_path = tmp.name
        plotter.screenshot(tmp_path)
        plotter.close()

        with open(tmp_path, 'rb') as f:
            img_bytes = f.read()
        os.unlink(tmp_path)
        return img_bytes

    except Exception as e:
        print(f"[レポート] スクリーンショット生成エラー: {e}")
        return None


# ============================================================================
# グラフ生成
# ============================================================================

def generate_timeseries_graph(
    frame_indices: List[int],
    values: List[float],
    ylabel: str,
    title: str,
    color: str = '#2E75B6',
    figsize: Tuple[float, float] = (8, 3.5),
) -> Optional[bytes]:
    """時系列グラフを生成してPNG画像のバイト列を返す"""
    if not _HAS_MPL:
        return None

    _setup_matplotlib_japanese()

    try:
        fig, ax = plt.subplots(figsize=figsize, dpi=150)
        ax.plot(frame_indices, values, color=color, linewidth=1.5, marker='o', markersize=2)
        ax.set_xlabel('フレーム番号', fontsize=10)
        ax.set_ylabel(ylabel, fontsize=10)
        ax.set_title(title, fontsize=12, fontweight='bold')
        ax.grid(True, alpha=0.3)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

        # 最大値にアノテーション
        if values:
            max_idx = np.argmax(values)
            ax.annotate(
                f'最大: {values[max_idx]:.3f}',
                xy=(frame_indices[max_idx], values[max_idx]),
                xytext=(10, 10), textcoords='offset points',
                fontsize=8, color='red',
                arrowprops=dict(arrowstyle='->', color='red', lw=0.8),
            )

        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    except Exception as e:
        print(f"[レポート] グラフ生成エラー: {e}")
        return None


def generate_multi_timeseries_graph(
    frame_indices: List[int],
    datasets: List[Dict[str, Any]],
    title: str,
    figsize: Tuple[float, float] = (8, 4),
) -> Optional[bytes]:
    """複数系列の時系列グラフ"""
    if not _HAS_MPL:
        return None

    _setup_matplotlib_japanese()

    try:
        fig, axes = plt.subplots(len(datasets), 1, figsize=figsize, dpi=150, sharex=True)
        if len(datasets) == 1:
            axes = [axes]

        colors = ['#2E75B6', '#C00000', '#548235', '#7030A0']
        for i, (ax, ds) in enumerate(zip(axes, datasets)):
            c = colors[i % len(colors)]
            ax.plot(frame_indices, ds['values'], color=c, linewidth=1.2, marker='o', markersize=1.5)
            ax.set_ylabel(ds.get('ylabel', ''), fontsize=9)
            ax.set_title(ds.get('label', ''), fontsize=10, fontweight='bold', loc='left')
            ax.grid(True, alpha=0.3)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

        axes[-1].set_xlabel('フレーム番号', fontsize=10)
        fig.suptitle(title, fontsize=12, fontweight='bold', y=1.02)
        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    except Exception as e:
        print(f"[レポート] 複合グラフ生成エラー: {e}")
        return None


# ============================================================================
# メイン: レポート生成関数
# ============================================================================

def generate_fem_report(
    output_path: str,
    # 解析結果サマリー
    fem_results: Optional[Any] = None,
    # 時系列データ（事前計算済みフレームデータ）
    frame_data: Optional[List[Dict[str, Any]]] = None,
    # ヒートマップメッシュ
    heatmap_mesh: Optional[Any] = None,
    dist_mesh: Optional[Any] = None,
    # パラメータ
    material_params: Optional[Dict[str, Any]] = None,
    contact_params: Optional[Dict[str, Any]] = None,
    mesh_info: Optional[Dict[str, Any]] = None,
    boundary_mode: str = "auto",
    # メタ情報
    title: str = "FEM接触解析レポート",
    subtitle: str = "",
    author: str = "",
    additional_notes: str = "",
) -> str:
    """FEM解析レポートをWord(.docx)形式で生成

    Args:
        output_path: 出力ファイルパス
        fem_results: FEMResults オブジェクト（単一フレームの結果）
        frame_data: フレームごとのデータリスト
            [{'frame': int, 'contact_area': float, 'peak_pressure': float,
              'n_contact_nodes': int, 'solve_time': float}, ...]
        heatmap_mesh: PyVistaメッシュ（スカラーデータ付き）
        dist_mesh: 遠位メッシュ（スクリーンショット用）
        material_params: {'E': float, 'nu': float, 'thickness': float, 'bone_E': float, 'bone_nu': float}
        contact_params: {'penalty_stiffness': float, 'contact_tolerance': float}
        mesh_info: {'prox_nodes': int, 'prox_elements': int, 'dist_nodes': int, ...}
        boundary_mode: 境界条件モード
        title: レポートタイトル
        subtitle: サブタイトル
        author: 作成者名
        additional_notes: 追加メモ

    Returns:
        出力ファイルパス
    """
    if not _HAS_DOCX:
        raise RuntimeError("python-docxが必要です。pip install python-docx")

    doc = Document()

    # ページ設定（A4）
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # スタイル設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # フッター（ページ番号）
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.style.font.size = Pt(8)
    fp.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # ページ番号フィールドを追加
    run = fp.add_run()
    fld_char_begin = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._r.append(fld_char_begin)
    run2 = fp.add_run()
    instr = run2._r.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instr.text = ' PAGE '
    run2._r.append(instr)
    run3 = fp.add_run()
    fld_char_end = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._r.append(fld_char_end)

    now = datetime.datetime.now()

    # ================================================================
    # タイトルページ
    # ================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 区切り線
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run('─' * 40)
    run_line.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_line.font.size = Pt(10)

    # メタ情報
    meta_items = [
        ('作成日時', now.strftime('%Y年%m月%d日 %H:%M')),
        ('解析ソフトウェア', 'FRS_Simulator v2.1 / FEM Contact Solver v2'),
    ]
    if author:
        meta_items.insert(0, ('作成者', author))

    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f'{label}: ')
        run_label.font.size = Pt(11)
        run_label.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_val = p.add_run(value)
        run_val.font.size = Pt(11)

    doc.add_page_break()

    # ================================================================
    # 目次（手動）
    # ================================================================
    doc.add_heading('目次', level=1)
    toc_items = ['1. 解析パラメータ', '2. 解析結果サマリー']
    if heatmap_mesh is not None:
        toc_items.append('3. 接触圧分布（ヒートマップ）')
    if frame_data and len(frame_data) > 1:
        toc_items.append(f'{len(toc_items) + 1}. 時系列解析結果')
    toc_items.append(f'{len(toc_items) + 1}. 備考')

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.0)

    doc.add_page_break()

    # ================================================================
    # 1. 解析パラメータ
    # ================================================================
    doc.add_heading('1. 解析パラメータ', level=1)

    # 1.1 材料特性
    doc.add_heading('1.1 材料特性', level=2)
    if material_params:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_style(table)

        hdr = table.rows[0].cells
        for i, text in enumerate(['パラメータ', '値', '単位']):
            hdr[i].text = text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            _set_cell_shading(hdr[i], 'D5E8F0')

        mat_rows = [
            ('軟骨 ヤング率 (E)', f"{material_params.get('E', 10.0):.1f}", 'MPa'),
            ('軟骨 ポアソン比 (ν)', f"{material_params.get('nu', 0.45):.2f}", '—'),
            ('軟骨 厚さ (t)', f"{material_params.get('thickness', 2.0):.1f}", 'mm'),
        ]
        if 'bone_E' in material_params:
            mat_rows.append(('骨 ヤング率 (E)', f"{material_params['bone_E']:.0f}", 'MPa'))
        if 'bone_nu' in material_params:
            mat_rows.append(('骨 ポアソン比 (ν)', f"{material_params['bone_nu']:.2f}", '—'))

        for param, val, unit in mat_rows:
            row = table.add_row().cells
            row[0].text = param
            row[1].text = val
            row[2].text = unit
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    else:
        doc.add_paragraph('材料パラメータ情報が提供されていません。', style='Normal')

    doc.add_paragraph()  # スペーサー

    # 1.2 接触パラメータ
    doc.add_heading('1.2 接触解析パラメータ', level=2)
    if contact_params:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_style(table)

        hdr = table.rows[0].cells
        for i, text in enumerate(['パラメータ', '値', '単位']):
            hdr[i].text = text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            _set_cell_shading(hdr[i], 'D5E8F0')

        ct_rows = [
            ('ペナルティ剛性', f"{contact_params.get('penalty_stiffness', 500.0):.1f}", 'MPa/mm'),
            ('接触判定閾値', f"{contact_params.get('contact_tolerance', 2.0):.1f}", 'mm'),
            ('境界条件モード', boundary_mode, '—'),
        ]
        for param, val, unit in ct_rows:
            row = table.add_row().cells
            row[0].text = param
            row[1].text = val
            row[2].text = unit
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    else:
        doc.add_paragraph('接触パラメータ情報が提供されていません。', style='Normal')

    doc.add_paragraph()

    # 1.3 メッシュ情報
    doc.add_heading('1.3 メッシュ情報', level=2)
    if mesh_info:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_style(table)

        hdr = table.rows[0].cells
        for i, text in enumerate(['項目', '値']):
            hdr[i].text = text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            _set_cell_shading(hdr[i], 'D5E8F0')

        for key, val in mesh_info.items():
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = str(val)
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    else:
        doc.add_paragraph('メッシュ情報が提供されていません。', style='Normal')

    doc.add_page_break()

    # ================================================================
    # 2. 解析結果サマリー
    # ================================================================
    doc.add_heading('2. 解析結果サマリー', level=1)

    if fem_results is not None:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_style(table)

        hdr = table.rows[0].cells
        for i, text in enumerate(['指標', '値']):
            hdr[i].text = text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            _set_cell_shading(hdr[i], 'E8D5C4')

        result_rows = [
            ('解析節点数', f"{fem_results.n_nodes:,}"),
            ('解析要素数', f"{fem_results.n_elements:,}"),
            ('接触節点数', f"{fem_results.n_contact_nodes:,}"),
            ('接触面積', f"{fem_results.contact_area:.2f} mm²"),
            ('全接触力', f"{fem_results.total_contact_force:.2f} N"),
            ('最大接触圧', f"{fem_results.peak_contact_pressure:.4f} MPa"),
        ]
        if len(fem_results.von_mises_stress) > 0:
            result_rows.append(('最大 von Mises 応力', f"{np.max(fem_results.von_mises_stress):.4f} MPa"))
        if len(fem_results.displacement_magnitude) > 0:
            result_rows.append(('最大変位', f"{np.max(fem_results.displacement_magnitude):.4f} mm"))
        result_rows.append(('解析時間', f"{fem_results.solve_time_sec:.2f} 秒"))

        for label, val in result_rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        # 結果の要約テキスト
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('解析概要: ').font.bold = True
        summary_text = (
            f"本解析では{fem_results.n_nodes:,}節点、{fem_results.n_elements:,}要素のメッシュに対して"
            f"ペナルティ法による接触解析を実施しました。"
            f"{fem_results.n_contact_nodes:,}個の接触節点が検出され、"
            f"接触面積は{fem_results.contact_area:.2f} mm²でした。"
            f"最大接触圧は{fem_results.peak_contact_pressure:.4f} MPaです。"
        )
        p.add_run(summary_text)
    else:
        doc.add_paragraph('FEM解析結果が提供されていません。', style='Normal')

    # ================================================================
    # 3. ヒートマップ（スクリーンショット）
    # ================================================================
    section_num = 3
    if heatmap_mesh is not None and _HAS_PV:
        doc.add_page_break()
        doc.add_heading(f'{section_num}. 接触圧分布（ヒートマップ）', level=1)

        scalars_to_capture = []
        if 'contact_pressure' in heatmap_mesh.array_names:
            scalars_to_capture.append(('contact_pressure', '接触圧分布', 'jet'))
        if 'von_mises_stress' in heatmap_mesh.array_names:
            scalars_to_capture.append(('von_mises_stress', 'von Mises応力分布', 'plasma'))
        if 'displacement_magnitude' in heatmap_mesh.array_names:
            scalars_to_capture.append(('displacement_magnitude', '変位量分布', 'viridis'))

        if not scalars_to_capture and heatmap_mesh.array_names:
            first_scalar = heatmap_mesh.array_names[0]
            scalars_to_capture.append((first_scalar, first_scalar, 'jet'))

        for scalar_name, caption, cmap in scalars_to_capture:
            doc.add_heading(caption, level=2)
            img_bytes = capture_heatmap_screenshot(
                heatmap_mesh, scalar_name=scalar_name, cmap=cmap,
                dist_mesh=dist_mesh, title=caption,
            )
            if img_bytes:
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Cm(15.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph(f'[{caption}のスクリーンショットを生成できませんでした]')
            doc.add_paragraph()

        section_num += 1

    # ================================================================
    # 4. 時系列グラフ
    # ================================================================
    if frame_data and len(frame_data) > 1:
        doc.add_page_break()
        doc.add_heading(f'{section_num}. 時系列解析結果', level=1)

        frames = [d.get('frame', i) for i, d in enumerate(frame_data)]

        # 接触圧の時系列
        peak_pressures = [d.get('peak_pressure', 0) for d in frame_data]
        if any(p > 0 for p in peak_pressures):
            doc.add_heading('最大接触圧の時間変化', level=2)
            img_bytes = generate_timeseries_graph(
                frames, peak_pressures,
                ylabel='最大接触圧 [MPa]',
                title='フレームごとの最大接触圧',
                color='#C00000',
            )
            if img_bytes:
                doc.add_picture(io.BytesIO(img_bytes), width=Cm(15.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # 接触面積の時系列
        contact_areas = [d.get('contact_area', 0) for d in frame_data]
        if any(a > 0 for a in contact_areas):
            doc.add_heading('接触面積の時間変化', level=2)
            img_bytes = generate_timeseries_graph(
                frames, contact_areas,
                ylabel='接触面積 [mm²]',
                title='フレームごとの接触面積',
                color='#2E75B6',
            )
            if img_bytes:
                doc.add_picture(io.BytesIO(img_bytes), width=Cm(15.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # 接触節点数の時系列
        n_contacts = [d.get('n_contact_nodes', 0) for d in frame_data]
        if any(n > 0 for n in n_contacts):
            doc.add_heading('接触節点数の時間変化', level=2)
            img_bytes = generate_timeseries_graph(
                frames, n_contacts,
                ylabel='接触節点数',
                title='フレームごとの接触節点数',
                color='#548235',
            )
            if img_bytes:
                doc.add_picture(io.BytesIO(img_bytes), width=Cm(15.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        # 統計サマリーテーブル
        doc.add_heading('時系列統計サマリー', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_style(table)

        hdr = table.rows[0].cells
        for i, text in enumerate(['指標', '最小値', '最大値', '平均値', '標準偏差']):
            hdr[i].text = text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
            _set_cell_shading(hdr[i], 'D5E8F0')

        stats_data = [
            ('最大接触圧 [MPa]', peak_pressures),
            ('接触面積 [mm²]', contact_areas),
            ('接触節点数', n_contacts),
        ]
        for label, values in stats_data:
            if any(v > 0 for v in values):
                arr = np.array(values)
                row = table.add_row().cells
                row[0].text = label
                row[1].text = f"{np.min(arr):.4f}"
                row[2].text = f"{np.max(arr):.4f}"
                row[3].text = f"{np.mean(arr):.4f}"
                row[4].text = f"{np.std(arr):.4f}"
                for cell in row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

        section_num += 1

    # ================================================================
    # 備考
    # ================================================================
    doc.add_page_break()
    doc.add_heading(f'{section_num}. 備考', level=1)

    notes = [
        '本解析はCST（Constant Strain Triangle）膜要素とペナルティ法による接触解析に基づいています。',
        '軟骨は線形弾性体（平面応力仮定）としてモデル化されています。',
        '骨は剛体として扱われています。',
        '接触検出には compute_implicit_distance による符号付き距離法を使用しています。',
    ]
    for note in notes:
        p = doc.add_paragraph(note, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    if additional_notes:
        doc.add_paragraph()
        doc.add_heading('追加メモ', level=2)
        doc.add_paragraph(additional_notes)

    # ================================================================
    # 保存
    # ================================================================
    doc.save(output_path)
    print(f"[レポート] Word文書を保存しました: {output_path}")
    return output_path


# ============================================================================
# シミュレータ連携用ヘルパー
# ============================================================================

def generate_report_from_simulator(
    app,
    output_path: Optional[str] = None,
    include_screenshot: bool = True,
    frame_data: Optional[List[Dict]] = None,
) -> Optional[str]:
    """FRS_Simulator のアプリケーションインスタンスからレポートを生成

    Args:
        app: MainMenuGUI インスタンス
        output_path: 出力パス（Noneの場合はダイアログで選択）
        include_screenshot: スクリーンショットを含めるか
        frame_data: 時系列データ

    Returns:
        生成されたファイルパス、またはNone
    """
    import tkinter as tk
    from tkinter import filedialog, messagebox

    if not _HAS_DOCX:
        messagebox.showerror("エラー", "python-docxが必要です。\npip install python-docx")
        return None

    # 出力パス
    if output_path is None:
        output_path = filedialog.asksaveasfilename(
            title="レポートの保存先を選択",
            defaultextension=".docx",
            filetypes=[("Word文書", "*.docx"), ("すべてのファイル", "*.*")],
            initialfile=f"FEM_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        )
        if not output_path:
            return None

    # パラメータ収集
    material_params = None
    contact_params = None
    try:
        material_params = {
            'E': app.fem_cart_E.get(),
            'nu': app.fem_cart_nu.get(),
            'thickness': app.fem_cart_thickness.get(),
            'bone_E': app.fem_bone_E.get(),
            'bone_nu': app.fem_bone_nu.get(),
        }
        contact_params = {
            'penalty_stiffness': app.fem_penalty_stiffness.get(),
            'contact_tolerance': app.fem_contact_tolerance.get(),
        }
    except Exception as e:
        print(f"[レポート] パラメータ取得エラー: {e}")

    # FEM結果
    fem_results = getattr(app, 'fem_last_results', None)

    # メッシュ情報
    mesh_info = {}
    if fem_results:
        mesh_info['解析節点数'] = f"{fem_results.n_nodes:,}"
        mesh_info['解析要素数'] = f"{fem_results.n_elements:,}"

    prox_cart = getattr(app, 'fem_last_prox_cart_mesh', None)
    dist_cart = getattr(app, 'fem_last_dist_cart_mesh', None)
    if prox_cart:
        mesh_info['近位軟骨メッシュ節点数'] = f"{prox_cart.n_points:,}"
    if dist_cart:
        mesh_info['遠位軟骨メッシュ節点数'] = f"{dist_cart.n_points:,}"

    # ヒートマップメッシュ
    heatmap_mesh = None
    if include_screenshot and fem_results and prox_cart:
        try:
            from fem_contact_solver_2 import apply_fem_results_to_mesh
            heatmap_mesh = prox_cart.copy()
            if fem_results.n_nodes == heatmap_mesh.n_points:
                apply_fem_results_to_mesh(heatmap_mesh, fem_results)
            else:
                heatmap_mesh = None
        except Exception as e:
            print(f"[レポート] ヒートマップメッシュ生成エラー: {e}")

    # 境界条件
    boundary_mode = "auto"
    try:
        boundary_mode = app.fem_boundary_mode.get()
    except Exception:
        pass

    # レポート生成
    try:
        result_path = generate_fem_report(
            output_path=output_path,
            fem_results=fem_results,
            frame_data=frame_data,
            heatmap_mesh=heatmap_mesh,
            dist_mesh=dist_cart if include_screenshot else None,
            material_params=material_params,
            contact_params=contact_params,
            mesh_info=mesh_info if mesh_info else None,
            boundary_mode=boundary_mode,
            title="FEM接触解析レポート",
            subtitle=f"生成日時: {datetime.datetime.now().strftime('%Y/%m/%d %H:%M')}",
        )
        messagebox.showinfo("完了", f"レポートを保存しました:\n{result_path}")
        return result_path
    except Exception as e:
        messagebox.showerror("エラー", f"レポート生成中にエラーが発生しました:\n{e}")
        import traceback
        traceback.print_exc()
        return None
